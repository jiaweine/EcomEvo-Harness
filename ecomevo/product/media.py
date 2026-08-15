from __future__ import annotations
import json, subprocess
from pathlib import Path
from typing import Any
from PIL import Image

DISPLAY_LIMIT=24_000
SEARCH_LIMIT=500_000


def _read_text(path:Path,limit:int=SEARCH_LIMIT)->str:
    with path.open('r',encoding='utf-8',errors='ignore') as fh:
        return fh.read(limit)


def _text_meta(meta:dict[str,Any],text:str)->dict[str,Any]:
    text=text[:SEARCH_LIMIT]
    meta.update({
        'chars':len(text),
        'lines':text.count('\n')+1 if text else 0,
        'text':text[:DISPLAY_LIMIT],
        'search_text':text,
        'preview':text[:1600],
        'search_truncated':len(text)>=SEARCH_LIMIT,
    })
    return meta


def probe_media(path:str|Path,mime:str)->dict[str,Any]:
    path=Path(path);meta={'kind':'file'}
    if mime.startswith('image/'):
        try:
            with Image.open(path) as im:meta.update({'kind':'image','width':im.width,'height':im.height,'mode':im.mode})
        except Exception:meta['kind']='image'
        return meta
    if mime.startswith(('video/','audio/')):
        kind='video' if mime.startswith('video/') else 'audio';meta['kind']=kind
        try:
            cp=subprocess.run(['ffprobe','-v','quiet','-print_format','json','-show_format','-show_streams',str(path)],capture_output=True,text=True,timeout=20)
            d=json.loads(cp.stdout or '{}');fmt=d.get('format',{});streams=d.get('streams',[]);meta.update({'duration':round(float(fmt.get('duration',0) or 0),2),'bit_rate':int(float(fmt.get('bit_rate',0) or 0))})
            for s in streams:
                if s.get('codec_type')=='video':meta.update({'width':s.get('width'),'height':s.get('height'),'fps':s.get('avg_frame_rate')})
                if s.get('codec_type')=='audio':meta.update({'sample_rate':s.get('sample_rate'),'channels':s.get('channels')})
        except Exception:pass
        return meta
    suffix=path.suffix.lower()
    if mime.startswith('text/') or suffix in {'.txt','.log','.json','.csv','.tsv','.yaml','.yml','.xml','.md'}:
        meta['kind']='text'
        try:_text_meta(meta,_read_text(path))
        except Exception:pass
        return meta
    if suffix=='.pdf' or mime=='application/pdf':
        meta['kind']='document'
        try:
            from pypdf import PdfReader
            r=PdfReader(str(path));chunks=[];chars=0;indexed=0
            for i,page in enumerate(r.pages[:1000]):
                value=(page.extract_text() or '').strip()
                indexed+=1
                if not value:continue
                piece=f'[[第{i+1}页]]\n{value}\n'
                remaining=SEARCH_LIMIT-chars
                if remaining<=0:break
                chunks.append(piece[:remaining]);chars+=min(len(piece),remaining)
                if chars>=SEARCH_LIMIT:break
            text=''.join(chunks)
            _text_meta(meta,text)
            meta.update({'pages':len(r.pages),'indexed_pages':indexed,'text_density':round(len(text)/max(1,indexed),1)})
        except Exception:pass
        return meta
    if suffix=='.docx':
        meta['kind']='document'
        try:
            from docx import Document
            doc=Document(str(path));chunks=[];chars=0
            for paragraph in doc.paragraphs:
                value=paragraph.text.strip()
                if value:
                    piece=value+'\n';remaining=SEARCH_LIMIT-chars
                    if remaining<=0:break
                    chunks.append(piece[:remaining]);chars+=min(len(piece),remaining)
            # Business documents often keep the important fields in tables rather than paragraphs.
            if chars<SEARCH_LIMIT:
                for table in doc.tables:
                    for row in table.rows:
                        piece='\t'.join(cell.text.strip() for cell in row.cells)+'\n';remaining=SEARCH_LIMIT-chars
                        if remaining<=0:break
                        chunks.append(piece[:remaining]);chars+=min(len(piece),remaining)
                    if chars>=SEARCH_LIMIT:break
            text=''.join(chunks);_text_meta(meta,text);meta.update({'paragraphs':len(doc.paragraphs),'tables':len(doc.tables)})
        except Exception:pass
        return meta
    if suffix in {'.xlsx','.xlsm'}:
        meta['kind']='sheet'
        try:
            from openpyxl import load_workbook
            wb=load_workbook(path,read_only=True,data_only=True);sn=wb.sheetnames[:8];chunks=[];rows=0;chars=0
            stop=False
            for name in sn:
                ws=wb[name];header=f'[[工作表:{name}]]\n';chunks.append(header);chars+=len(header)
                for row in ws.iter_rows(min_row=1,max_row=5000,values_only=True):
                    piece='\t'.join('' if v is None else str(v) for v in row)+'\n';remaining=SEARCH_LIMIT-chars
                    if remaining<=0:stop=True;break
                    chunks.append(piece[:remaining]);chars+=min(len(piece),remaining);rows+=1
                    if chars>=SEARCH_LIMIT:stop=True;break
                if stop:break
            text=''.join(chunks);_text_meta(meta,text);meta.update({'sheets':sn,'rows':rows,'indexed_rows':rows})
        except Exception:pass
        return meta
    return meta


def extract_video_frames(path:str|Path,out_dir:str|Path,count:int=4)->list[str]:
    path=Path(path);out_dir=Path(out_dir);out_dir.mkdir(parents=True,exist_ok=True)
    try:
        probe=probe_media(path,'video/mp4');dur=float(probe.get('duration',0) or 0)
        if dur<=0:return []
        times=[dur*(i+1)/(count+1) for i in range(count)];rows=[]
        for i,t in enumerate(times,1):
            dest=out_dir/f'frame_{i}.jpg';subprocess.run(['ffmpeg','-y','-ss',str(t),'-i',str(path),'-frames:v','1','-vf',"scale='min(1280,iw)':-2",str(dest)],stdout=subprocess.DEVNULL,stderr=subprocess.DEVNULL,timeout=30)
            if dest.exists():rows.append(str(dest))
        return rows
    except Exception:return []
